import tiktoken
from typing import List
import io
import uuid
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from src.chatup.db.pinecone import upsert_embeddings
from src.chatup.config import settings

# Initialize embeddings model
embeddings_model = OpenAIEmbeddings(openai_api_key=settings.OPENAI_API_KEY)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from different file types given as bytes.
    Supports: .txt, .pdf, .docx, .png/.jpg/.jpeg (via OCR)
    """
    extension = filename.split('.')[-1].lower()

    if extension == "txt":
        return file_bytes.decode("utf-8")

    elif extension == "pdf":
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    # Keep track of page number as well
                    text += f"[Page {i}]\n{page_text}\n"
                tables = page.extract_tables()
                for table in tables:
                    # Convert each table (list of rows) to a string
                    table_str = "\n".join([" | ".join(cell if cell else "" for cell in row) for row in table])
                    text += f"[Page {i} - Table]\n{table_str}\n"
        return text
    elif extension == "docx":
        text = ""
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                text += " | ".join([cell.text for cell in row.cells]) + "\n"
        return text

    elif extension in ["png", "jpg", "jpeg"]:
        image = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(image)

    else:
        raise ValueError(f"Unsupported file type: {extension}")


def smart_split_text(text: str, max_tokens: int = 2000, model_name="text-embedding-ada-002"):
    """
    Smart recursive splitter:
    - Prefers splitting by paragraph, sentence, newline.
    - Falls back to characters if needed.
    - Respects max_tokens limit based on the embedding model.
    """
    enc = tiktoken.encoding_for_model(model_name)
    
    def count_tokens(s: str) -> int:
        return len(enc.encode(s))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_tokens,
        chunk_overlap=50,
        length_function=count_tokens,
        separators=["\n\n", "\n", ".", "?", "!", " "]  # priority of splits
    )
    return splitter.split_text(text)

def semantic_split_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    """
    Semantically split text using embeddings similarity + optional chunk size control.
    Batches embedding requests to avoid OpenAI token limits (by token count).
    Splits very large paragraphs by lines if needed.
    """
    paragraphs = smart_split_text(text)
    
    paragraph_embeddings = []
    for para_batch in paragraphs:
        paragraph_embeddings.extend(embeddings_model.embed_documents(para_batch))

    chunks = []
    current_chunk = ""
    current_embs = []

    for para, emb in zip(paragraphs, paragraph_embeddings):
        current_chunk += para + "\n\n"
        current_embs.append(emb)
        
        avg_emb = np.mean(current_embs, axis=0)
        similarity = np.dot(emb, avg_emb) / (np.linalg.norm(emb) * np.linalg.norm(avg_emb))

        if len(current_chunk) > chunk_size or similarity < 0.95:
            chunks.append(current_chunk.strip())
            current_chunk = ""
            current_embs = []

    if current_chunk:
        chunks.append(current_chunk.strip())
    
    # Add overlap
    final_chunks = []
    for i, c in enumerate(chunks):
        final_chunks.append(c)
        if i < len(chunks) - 1:
            overlap = chunks[i + 1][:chunk_overlap]
            final_chunks[-1] += "\n" + overlap

    return final_chunks


def embed_and_upload(chunks: List[str], namespace: str = "default", source: str = "unknown"):
    """
    Create embeddings for each chunk and upload to Pinecone under the given namespace.
    Store vector, text, and source metadata (filename, page info).
    """
    vectors = []
    for i, chunk in enumerate(chunks):
        vector = embeddings_model.embed_query(chunk)
        vectors.append({
            "id": f"chunk-{str(uuid.uuid4())[:8]}",
            "values": vector,
            "metadata": {
                "text": chunk,
                "source": source,
                "namespace": namespace
            }
        })

    upsert_embeddings(vectors=vectors, namespace=namespace)
